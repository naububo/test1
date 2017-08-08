from docx import Document
from log import logger



def read_sheets_from_template(template_file):
    wordDoc = Document(template_file)
    tables = wordDoc.tables
    sheets = []

    for t_idx, table in enumerate(tables):
        sheet = []

        for r_idx, row in enumerate(table.rows):
            row_data = []
            for c_idx, cell in enumerate(row.cells):
                if r_idx == 0:
                    if len(cell.text) >= 0:
                        row_data.append(cell.text)
                else:
                    row_data.append(cell.text)
            if len(row_data) > 0:
                sheet.append(row_data)
        if len(sheet) > 0:
            sheets.append(sheet)
    print(sheets)
    return sheets


def write_to_docx(sheets,docx_template):
    wordDoc = Document(docx_template)
    # para = wordDoc.add_paragraph()
    tables = wordDoc.tables
    sStr2 = "Tr"
    for t_idx, table in enumerate(tables):
        if table is not None:
            for r_idx, row in enumerate(table.rows):
                if row is not None:
                    for c_idx, cell in enumerate(row.cells):
                        if r_idx >= 0 and c_idx >= 0 and len(cell.text) >= 0:
                            try:
                                if cell is not None:
                                    cell.text = str(sheets[t_idx][r_idx][c_idx])
                                    # <----remove flag string --->
                                    sStr1 =cell.text
                                    # print(sStr1)
                                    if sStr2 in sStr1:
                                        nPos = sStr1.index(sStr2)
                                        cell.text=sStr1[0:nPos]
                                        # print (nPos)
                                    # <--------------------------->
                            except Exception as err:
                                logger.error(
                                    "t_idx  " + str(t_idx) + "   r_idx  " + str(r_idx) + "    c_idx  " + str(c_idx))
                                logger.error(err)

    wordDoc.save(docx_template)

if __name__ == '__main__':

    sheets1 = read_sheets_from_template("../docx/test1.docx")
    write_to_docx(sheets1,"../docx/test1.docx")
